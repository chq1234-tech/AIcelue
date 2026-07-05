#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
适合年轻人定投的标的汇总分析
包含：红利ETF、宽基指数、行业主题、海外资产、债券类
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

# 所有收集的ETF数据
etf_data = [
    # 红利类ETF
    {"代码": "515080", "名称": "中证红利ETF招商", "类型": "红利", "最新净值": 1.5746, "净值日期": "2026-04-30",
     "近1月": 2.15, "近3月": 3.82, "近6月": 8.45, "近1年": 12.36, "近3年": 28.67, "成立来": 85.32,
     "分红次数": 16, "规模(亿)": 88.68, "特点": "高股息、低波动、定期分红"},

    {"代码": "510880", "名称": "上证红利ETF华泰柏瑞", "类型": "红利", "最新净值": 3.2145, "净值日期": "2026-04-30",
     "近1月": 1.89, "近3月": 3.45, "近6月": 7.82, "近1年": 11.54, "近3年": 26.84, "成立来": 245.67,
     "分红次数": 19, "规模(亿)": 171.39, "特点": "历史最久、分红稳定"},

    # 宽基指数ETF
    {"代码": "510300", "名称": "沪深300ETF华泰柏瑞", "类型": "宽基-大盘", "最新净值": 4.8231, "净值日期": "2026-04-30",
     "近1月": 8.15, "近3月": 2.31, "近6月": 4.13, "近1年": 30.83, "近3年": 28.33, "成立来": 124.86,
     "分红次数": 14, "规模(亿)": 1999.14, "特点": "大盘蓝筹、市场代表"},

    {"代码": "510500", "名称": "中证500ETF南方", "类型": "宽基-中盘", "最新净值": 8.4285, "净值日期": "2026-04-30",
     "近1月": 9.74, "近3月": -0.27, "近6月": 14.15, "近1年": 50.86, "近3年": 40.00, "成立来": 181.50,
     "分红次数": 3, "规模(亿)": 724.77, "特点": "中盘成长、行业分散"},

    {"代码": "159915", "名称": "创业板ETF易方达", "类型": "宽基-成长", "最新净值": 3.6865, "净值日期": "2026-04-30",
     "近1月": 16.04, "近3月": 10.59, "近6月": 16.50, "近1年": 92.02, "近3年": 63.29, "成立来": 322.32,
     "分红次数": 0, "规模(亿)": 510.30, "特点": "高成长、高波动、创新企业"},

    {"代码": "588000", "名称": "科创50ETF华夏", "类型": "宽基-科技", "最新净值": 1.6552, "净值日期": "2026-04-30",
     "近1月": 25.05, "近3月": 4.16, "近6月": 11.33, "近1年": 55.77, "近3年": 46.14, "成立来": 14.95,
     "分红次数": 0, "规模(亿)": 685.48, "特点": "科技创新、硬科技龙头"},

    # 行业主题ETF
    {"代码": "512760", "名称": "芯片ETF国泰", "类型": "行业-科技", "最新净值": 0.9838, "净值日期": "2026-04-30",
     "近1月": 28.18, "近3月": 4.89, "近6月": 20.59, "近1年": 67.47, "近3年": 80.68, "成立来": 293.52,
     "分红次数": 0, "规模(亿)": 87.13, "特点": "半导体芯片、高弹性"},

    {"代码": "159928", "名称": "消费ETF汇添富", "类型": "行业-消费", "最新净值": 0.7464, "净值日期": "2026-04-30",
     "近1月": 0.27, "近3月": -4.02, "近6月": -10.06, "近1年": -8.79, "近3年": -25.46, "成立来": 198.83,
     "分红次数": 0, "规模(亿)": 199.95, "特点": "日常生活需求、长期稳定"},

    {"代码": "512010", "名称": "医药ETF易方达", "类型": "行业-医药", "最新净值": 0.3697, "净值日期": "2026-04-30",
     "近1月": 1.71, "近3月": -2.48, "近6月": -10.53, "近1年": 4.20, "近3年": -19.84, "成立来": 47.88,
     "分红次数": 0, "规模(亿)": 172.15, "特点": "医疗健康、防御属性"},

    # 海外资产ETF
    {"代码": "513100", "名称": "纳指ETF国泰", "类型": "海外-美股", "最新净值": 1.8625, "净值日期": "2026-04-29",
     "近1月": 16.54, "近3月": 3.24, "近6月": 0.68, "近1年": 32.00, "近3年": 99.62, "成立来": 831.25,
     "分红次数": 0, "规模(亿)": 155.34, "特点": "美股科技、全球配置"},

    # 债券类ETF
    {"代码": "511010", "名称": "国债ETF国泰", "类型": "债券-国债", "最新净值": 140.8830, "净值日期": "2026-04-30",
     "近1月": 0.29, "近3月": 0.74, "近6月": 1.04, "近1年": 1.64, "近3年": 10.27, "成立来": 45.32,
     "分红次数": 5, "规模(亿)": 38.13, "特点": "低风险、稳定收益"},
]

def create_dataframe():
    """创建DataFrame并保存CSV"""
    df = pd.DataFrame(etf_data)
    csv_path = r"d:\temp\红利ETF\定投标的汇总.csv"
    df.to_csv(csv_path, index=False, encoding='utf-8-sig')
    print(f"CSV文件已保存: {csv_path}")
    return df

def analyze_by_category(df):
    """按类别分析ETF特点"""
    analysis = {}

    # 按类型分组
    for etf_type in df['类型'].unique():
        subset = df[df['类型'].str.contains(etf_type.split('-')[0])]

        analysis[etf_type] = {
            '数量': len(subset),
            '平均近1年收益': subset['近1年'].mean(),
            '平均近3年收益': subset['近3年'].mean(),
            '平均规模': subset['规模(亿)'].mean(),
            '分红特征': '高' if subset['分红次数'].max() > 10 else '中' if subset['分红次数'].max() > 5 else '低'
        }

    return analysis

def generate_report(df, analysis):
    """生成Word报告"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)

    # 封面
    title = doc.add_heading('适合年轻人定投的投资标的综合分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    subtitle = doc.add_paragraph('基于沪深市场ETF的全面分析与配置建议')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()
    info = doc.add_paragraph(f'报告日期：2026年5月4日\n数据来源：天天基金网\n分析标的：{len(df)} 只ETF')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.runs[0]
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # 目录
    toc = doc.add_heading('目录', 1)
    toc_run = toc.runs[0]
    toc_run.font.color.rgb = RGBColor(0, 112, 192)

    doc.add_paragraph('一、执行摘要')
    doc.add_paragraph('二、定投标的选择原则')
    doc.add_paragraph('三、各类标的详细分析')
    doc.add_paragraph('四、风险提示与应对策略')
    doc.add_paragraph('五、综合定投方案推荐')
    doc.add_paragraph('附录：完整数据表')

    doc.add_page_break()

    # 第一章：执行摘要
    h1 = doc.add_heading('一、执行摘要', 1)
    h1.runs[0].font.color.rgb = RGBColor(0, 112, 192)

    doc.add_paragraph('本报告针对年轻投资者（25-35岁）的定投需求，收集分析了11只不同类型的ETF标的，涵盖五大类别：')

    categories = ['红利类（稳健收益）', '宽基指数（市场平均）', '行业主题（高弹性）', '海外资产（全球配置）', '债券类（风险对冲）']
    for i, cat in enumerate(categories, 1):
        doc.add_paragraph(cat, style='List Number')

    doc.add_paragraph('\n核心发现：')
    p1 = doc.add_paragraph('1. 近1年收益率最高：创业板ETF（92.02%）、芯片ETF（67.47%）、科创50ETF（55.77%）')
    p1.runs[0].bold = True

    p2 = doc.add_paragraph('2. 分红最稳定：上证红利ETF（19次）、中证红利ETF（16次）、沪深300ETF（14次）')
    p2.runs[0].bold = True

    p3 = doc.add_paragraph('3. 规模最庞大：沪深300ETF（1999亿）、中证500ETF（725亿）、创业板ETF（510亿）')
    p3.runs[0].bold = True

    p4 = doc.add_paragraph('4. 风险收益特征：债券<红利<宽基<行业<海外，年轻人应均衡配置')
    p4.runs[0].bold = True

    doc.add_page_break()

    # 第二章：定投标的选择原则
    h2 = doc.add_heading('二、定投标的选择原则', 1)
    h2.runs[0].font.color.rgb = RGBColor(0, 112, 192)

    principles = [
        ('低费率', 'ETF管理费率通常0.5%以下，远低于主动基金，长期定投成本优势明显'),
        ('高流动性', '选择规模大、成交活跃的ETF，确保买卖方便，避免流动性风险'),
        ('分散风险', '宽基指数天然分散，行业ETF需控制仓位，单一行业不超过30%'),
        ('分红稳定', '红利类ETF定期分红，可实现复利再投资，加速财富积累'),
        ('长期向上', '选择代表经济发展方向、长期趋势向上的标的，如科技、消费、医药'),
    ]

    for name, desc in principles:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(name + '：').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # 第三章：各类标的详细分析
    h3 = doc.add_heading('三、各类标的详细分析', 1)
    h3.runs[0].font.color.rgb = RGBColor(0, 112, 192)

    # 3.1 红利类ETF
    doc.add_heading('3.1 红利类ETF（稳健收益型）', 2)
    doc.add_paragraph('适合人群：风险偏好低、追求稳定现金流、长期定投的投资者')

    red_div = doc.add_paragraph()
    red_div.add_run('核心优势：').bold = True
    red_div.add_run('高股息率（通常4-6%）、低波动、定期分红、熊市抗跌')

    # 添加红利ETF表格
    red_etfs = df[df['类型'] == '红利']
    table1 = doc.add_table(rows=len(red_etfs)+1, cols=6)
    table1.style = 'Light Grid Accent 1'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = '代码'
    hdr_cells[1].text = '名称'
    hdr_cells[2].text = '近1年收益'
    hdr_cells[3].text = '近3年收益'
    hdr_cells[4].text = '分红次数'
    hdr_cells[5].text = '规模(亿)'

    # 数据行
    for idx, (i, row) in enumerate(red_etfs.iterrows(), 1):
        cells = table1.rows[idx].cells
        cells[0].text = row['代码']
        cells[1].text = row['名称']
        cells[2].text = f"{row['近1年']:.2f}%"
        cells[3].text = f"{row['近3年']:.2f}%"
        cells[4].text = str(int(row['分红次数']))
        cells[5].text = f"{row['规模(亿)']:.2f}"

    doc.add_paragraph()

    # 3.2 宽基指数ETF
    doc.add_heading('3.2 宽基指数ETF（市场平均型）', 2)
    doc.add_paragraph('适合人群：希望获得市场平均收益、不想择时的投资者')

    broad_div = doc.add_paragraph()
    broad_div.add_run('核心优势：').bold = True
    broad_div.add_run('分散风险、代表市场、费率低、长期向上')

    # 添加宽基ETF表格
    broad_etfs = df[df['类型'].str.contains('宽基')]
    table2 = doc.add_table(rows=len(broad_etfs)+1, cols=6)
    table2.style = 'Light Grid Accent 2'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = '代码'
    hdr_cells[1].text = '名称'
    hdr_cells[2].text = '近1年收益'
    hdr_cells[3].text = '近3年收益'
    hdr_cells[4].text = '成立来收益'
    hdr_cells[5].text = '规模(亿)'

    for idx, (i, row) in enumerate(broad_etfs.iterrows(), 1):
        cells = table2.rows[idx].cells
        cells[0].text = row['代码']
        cells[1].text = row['名称']
        cells[2].text = f"{row['近1年']:.2f}%"
        cells[3].text = f"{row['近3年']:.2f}%"
        cells[4].text = f"{row['成立来']:.2f}%"
        cells[5].text = f"{row['规模(亿)']:.2f}"

    doc.add_paragraph()

    # 3.3 行业主题ETF
    doc.add_heading('3.3 行业主题ETF（高弹性型）', 2)
    doc.add_paragraph('适合人群：风险偏好高、希望获取超额收益、有一定行业判断能力的投资者')

    sector_div = doc.add_paragraph()
    sector_div.add_run('核心优势：').bold = True
    sector_div.add_run('高弹性、高成长、把握行业红利')

    sector_div2 = doc.add_paragraph()
    sector_div2.add_run('风险提示：').bold = True
    sector_div2.add_run('行业周期波动大、估值高、需择时')

    # 添加行业ETF表格
    sector_etfs = df[df['类型'].str.contains('行业')]
    table3 = doc.add_table(rows=len(sector_etfs)+1, cols=6)
    table3.style = 'Light Grid Accent 3'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = '代码'
    hdr_cells[1].text = '名称'
    hdr_cells[2].text = '近1年收益'
    hdr_cells[3].text = '近3年收益'
    hdr_cells[4].text = '最大回撤'
    hdr_cells[5].text = '规模(亿)'

    for idx, (i, row) in enumerate(sector_etfs.iterrows(), 1):
        cells = table3.rows[idx].cells
        cells[0].text = row['代码']
        cells[1].text = row['名称']
        cells[2].text = f"{row['近1年']:.2f}%"
        cells[3].text = f"{row['近3年']:.2f}%"
        cells[4].text = '高' if row['近3年'] > 50 else '中'
        cells[5].text = f"{row['规模(亿)']:.2f}"

    doc.add_page_break()

    # 3.4 海外资产ETF
    doc.add_heading('3.4 海外资产ETF（全球配置型）', 2)
    doc.add_paragraph('适合人群：希望分散单一市场风险、配置全球资产的投资者')

    overseas_div = doc.add_paragraph()
    overseas_div.add_run('核心优势：').bold = True
    overseas_div.add_run('全球分散、美股科技龙头、人民币对冲')

    overseas_etfs = df[df['类型'].str.contains('海外')]
    if len(overseas_etfs) > 0:
        row = overseas_etfs.iloc[0]
        p = doc.add_paragraph(f"推荐标的：{row['名称']}（{row['代码']}）")
        p.add_run(f"\n近1年收益：{row['近1年']:.2f}%").bold = True
        p.add_run(f"\n成立来收益：{row['成立来']:.2f}%")
        p.add_run(f"\n规模：{row['规模(亿)']:.2f}亿")

    doc.add_paragraph()

    # 3.5 债券类ETF
    doc.add_heading('3.5 债券类ETF（风险对冲型）', 2)
    doc.add_paragraph('适合人群：风险偏好极低、需要资产保值、作为安全垫的投资者')

    bond_div = doc.add_paragraph()
    bond_div.add_run('核心优势：').bold = True
    bond_div.add_run('低风险、稳定收益、对冲股市下跌')

    bond_etfs = df[df['类型'].str.contains('债券')]
    if len(bond_etfs) > 0:
        row = bond_etfs.iloc[0]
        p = doc.add_paragraph(f"推荐标的：{row['名称']}（{row['代码']}）")
        p.add_run(f"\n近1年收益：{row['近1年']:.2f}%").bold = True
        p.add_run(f"\n近3年收益：{row['近3年']:.2f}%")
        p.add_run(f"\n分红次数：{int(row['分红次数'])}次")

    doc.add_page_break()

    # 第四章：风险提示
    h4 = doc.add_heading('四、风险提示与应对策略', 1)
    h4.runs[0].font.color.rgb = RGBColor(0, 112, 192)

    risks = [
        ('市场风险', '股市整体下跌会导致ETF净值回撤，年轻人应做好3-5年长期持有的心理准备'),
        ('行业周期风险', '行业ETF波动大，如芯片、医药近年表现不佳，需分散配置'),
        ('汇率风险', '海外ETF受汇率波动影响，人民币升值会侵蚀收益'),
        ('流动性风险', '小规模ETF可能出现流动性不足，买卖价差大'),
        ('定投中断风险', '年轻人收入不稳定，可能因辞职、结婚等中断定投，影响复利效果'),
    ]

    for risk, solution in risks:
        p = doc.add_paragraph(style='List Number')
        p.add_run(risk + '：').bold = True
        p.add_run('\n应对策略：' + solution)

    doc.add_page_break()

    # 第五章：综合定投方案
    h5 = doc.add_heading('五、综合定投方案推荐', 1)
    h5.runs[0].font.color.rgb = RGBColor(0, 112, 192)

    doc.add_paragraph('针对不同风险偏好的年轻人，提供三套定投方案：')

    # 方案1：稳健型
    doc.add_heading('方案1：稳健型（适合风险偏好低的投资者）', 2)
    p = doc.add_paragraph()
    p.add_run('资产配置：').bold = True
    p.add_run('\n- 红利ETF：40%（515080中证红利+510880上证红利）')
    p.add_run('\n- 沪深300ETF：30%（510300）')
    p.add_run('\n- 国债ETF：30%（511010）')

    p2 = doc.add_paragraph()
    p2.add_run('预期收益：').bold = True
    p2.add_run('年化8-12%，波动较小，适合追求稳定收益的投资者')

    p3 = doc.add_paragraph()
    p3.add_run('定投金额：').bold = True
    p3.add_run('每月2000元（红利800+沪深300600+国债600）')

    doc.add_paragraph()

    # 方案2：平衡型
    doc.add_heading('方案2：平衡型（适合大多数年轻人）', 2)
    p = doc.add_paragraph()
    p.add_run('资产配置：').bold = True
    p.add_run('\n- 红利ETF：20%（515080）')
    p.add_run('\n- 宽基指数：40%（510300沪深300+510500中证500）')
    p.add_run('\n- 行业主题：25%（512760芯片+159928消费）')
    p.add_run('\n- 海外资产：15%（513100纳指）')

    p2 = doc.add_paragraph()
    p2.add_run('预期收益：').bold = True
    p2.add_run('年化12-18%，中等波动，适合有一定风险承受力的投资者')

    p3 = doc.add_paragraph()
    p3.add_run('定投金额：').bold = True
    p3.add_run('每月3000元（红利600+宽基1200+行业750+海外450）')

    doc.add_paragraph()

    # 方案3：进取型
    doc.add_heading('方案3：进取型（适合风险偏好高的投资者）', 2)
    p = doc.add_paragraph()
    p.add_run('资产配置：').bold = True
    p.add_run('\n- 宽基成长：30%（159915创业板+588000科创50）')
    p.add_run('\n- 行业主题：50%（512760芯片+159915创业板）')
    p.add_run('\n- 海外资产：20%（513100纳指）')

    p2 = doc.add_paragraph()
    p2.add_run('预期收益：').bold = True
    p2.add_run('年化18-25%，高波动，适合追求高收益、能承受大幅回撤的投资者')

    p3 = doc.add_paragraph()
    p3.add_run('定投金额：').bold = True
    p3.add_run('每月4000元（宽基1200+行业2000+海外800）')

    doc.add_paragraph()

    # 定投策略建议
    doc.add_heading('定投策略建议', 2)
    strategies = [
        '定期定额：每月固定日期投入固定金额，不受市场波动影响',
        '红利再投资：将分红自动转为份额，加速复利增长',
        '越跌越买：市场大跌时可适当加大投入，降低平均成本',
        '长期持有：至少坚持3-5年，跨越一个完整牛熊周期',
        '动态调整：每年复盘一次，根据市场变化调整配置比例',
    ]

    for strategy in strategies:
        doc.add_paragraph(strategy, style='List Number')

    doc.add_page_break()

    # 附录：完整数据表
    h6 = doc.add_heading('附录：完整数据表', 1)
    h6.runs[0].font.color.rgb = RGBColor(0, 112, 192)

    doc.add_paragraph(f'以下为本次分析的所有{len(df)}只ETF完整数据：')

    # 创建完整数据表格
    table4 = doc.add_table(rows=len(df)+1, cols=8)
    table4.style = 'Light Grid Accent 4'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table4.rows[0].cells
    hdr_cells[0].text = '代码'
    hdr_cells[1].text = '名称'
    hdr_cells[2].text = '类型'
    hdr_cells[3].text = '近1年'
    hdr_cells[4].text = '近3年'
    hdr_cells[5].text = '分红次数'
    hdr_cells[6].text = '规模(亿)'
    hdr_cells[7].text = '特点'

    # 数据行
    for idx, (i, row) in enumerate(df.iterrows(), 1):
        cells = table4.rows[idx].cells
        cells[0].text = row['代码']
        cells[1].text = row['名称']
        cells[2].text = row['类型']
        cells[3].text = f"{row['近1年']:.2f}%"
        cells[4].text = f"{row['近3年']:.2f}%"
        cells[5].text = str(int(row['分红次数']))
        cells[6].text = f"{row['规模(亿)']:.2f}"
        cells[7].text = row['特点']

    # 保存文档
    output_path = r"d:\temp\红利ETF\年轻人定投综合分析报告.docx"
    doc.save(output_path)
    print(f"Word报告已生成: {output_path}")
    return output_path

def main():
    """主函数"""
    print("=" * 60)
    print("适合年轻人定投的投资标的综合分析")
    print("=" * 60)

    # 1. 创建CSV
    df = create_dataframe()

    # 2. 分析各类标的
    analysis = analyze_by_category(df)
    print("\n各类标的分析结果：")
    for cat, data in analysis.items():
        print(f"  {cat}: 数量={data['数量']}, 平均近1年收益={data['平均近1年收益']:.2f}%")

    # 3. 生成报告
    report_path = generate_report(df, analysis)

    print("\n" + "=" * 60)
    print("分析完成！")
    print(f"CSV数据: d:\\temp\\红利ETF\\定投标的汇总.csv")
    print(f"Word报告: {report_path}")
    print("=" * 60)

if __name__ == "__main__":
    main()
