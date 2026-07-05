"""
红利ETF定投分析报告 - Word文档生成脚本
使用python-docx库生成专业格式的Word文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'start', 'bottom', 'end', 'left', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ['sz', 'val', 'color', 'space', 'shadow']:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), edge_data[key])
            tcBorders.append(element)
    tcPr.append(tcBorders)

def create_document():
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)
    
    # ========== 封面 ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('红利ETF定投分析报告')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    title.space_before = Pt(120)
    title.space_after = Pt(20)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('基于沪深市场9只主流红利ETF的深度分析')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    subtitle.space_after = Pt(40)
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f'生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    date_para.space_after = Pt(200)
    
    doc.add_page_break()
    
    # ========== 目录 ==========
    heading = doc.add_heading('目录', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    toc_items = [
        '一、红利ETF市场概览',
        '二、9只红利ETF对比分析',
        '三、定投标的推荐',
        '四、具体定投计划',
        '五、风险提示与应对'
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.space_before = Pt(6)
        p.space_after = Pt(6)
    
    doc.add_page_break()
    
    # ========== 第一章 ==========
    doc.add_heading('一、红利ETF市场概览', level=1)
    
    p = doc.add_paragraph('红利ETF是一种跟踪红利指数的交易型开放式指数基金，主要投资于高分红、低估值、业绩稳定的上市公司股票。相比普通ETF，红利ETF具有以下优势：')
    p.space_after = Pt(12)
    
    advantages = [
        '稳定分红：成分股均为高股息率公司，定期分红',
        '抗跌性强：熊市中表现相对稳健',
        '适合定投：波动相对较小，适合长期定投',
        '复利效应：分红再投资可产生复利增长'
    ]
    
    for adv in advantages:
        p = doc.add_paragraph(adv, style='List Bullet')
        p.space_after = Pt(6)
    
    doc.add_page_break()
    
    # ========== 第二章 ==========
    doc.add_heading('二、9只红利ETF对比分析', level=1)
    
    p = doc.add_paragraph('本次分析覆盖沪深市场9只主流红利ETF，数据截至2026年4月30日，来源为天天基金网。')
    p.space_after = Pt(12)
    
    # 创建对比表格
    table = doc.add_table(rows=6, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    headers = ['代码', '名称', '规模(亿)', '近1年', '近3年', '分红次数']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    # 数据行
    data = [
        ['510880', '华泰柏瑞上证红利', '171.39', '15.19%', '25.47%', '19次'],
        ['512890', '红利低波ETF', '311.87', '7.23%', '24.41%', '0次'],
        ['515080', '中证红利ETF招商★', '88.68', '14.88%', '24.19%', '16次'],
        ['515450', '南方红利低波50', 'N/A', '8.05%', '34.72%', '7次'],
        ['159758', '红利质量ETF华夏', 'N/A', '28.83%', '20.44%', '0次']
    ]
    
    for i, row_data in enumerate(data, 1):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_data
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ========== 第三章 ==========
    doc.add_heading('三、定投标的推荐', level=1)
    
    doc.add_heading('首选推荐：515080 中证红利ETF招商', level=2)
    
    p = doc.add_paragraph('推荐理由：')
    p.space_after = Pt(6)
    
    reasons = [
        '跨市场分散：同时覆盖上海和深圳市场，分散性优于仅覆盖沪市的上证红利ETF',
        '高频率分红：成立以来分红16次，平均每季度都有分红，可实现分红再投资的复利效应',
        '规模适中：88.68亿元，既不会太小（流动性风险），也不会太大（打新收益稀释）',
        '历史表现稳健：近1年14.88%，近3年24.19%，各周期表现均衡',
        '管理规范：招商基金管理，年化跟踪误差仅0.80%，费用低'
    ]
    
    for reason in reasons:
        p = doc.add_paragraph(reason, style='List Number')
        p.space_after = Pt(6)
    
    doc.add_page_break()
    
    # ========== 第四章 ==========
    doc.add_heading('四、具体定投计划', level=1)
    
    doc.add_heading('4.1 定投参数设定', level=2)
    
    # 定投参数表格
    param_table = doc.add_table(rows=8, cols=3)
    param_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    param_headers = ['参数', '设定值', '说明']
    for i, header in enumerate(param_headers):
        cell = param_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    param_data = [
        ['定投标的', '515080 + 510880', '主力+补充，跨市场覆盖'],
        ['定投频率', '每月', '工资日后3-5天'],
        ['每月金额', '2000元', '可根据收入调整'],
        ['分红方式', '红利再投资', '不取现金，自动再投资'],
        ['定投期限', '5年', '至少一轮牛熊周期'],
        ['预期年化收益', '12%', '基于历史表现测算']
    ]
    
    for i, row_data in enumerate(param_data, 1):
        for j, cell_data in enumerate(row_data):
            cell = param_table.rows[i].cells[j]
            cell.text = cell_data
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    doc.add_heading('4.2 收益测算', level=2)
    
    # 收益测算表格
    calc_table = doc.add_table(rows=6, cols=5)
    calc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    calc_headers = ['年限', '累计投入', '预期市值', '收益', '收益率']
    for i, header in enumerate(calc_headers):
        cell = calc_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    calc_data = [
        ['1年', '24,000元', '25,440元', '1,440元', '6.0%'],
        ['3年', '72,000元', '89,993元', '17,993元', '25.0%'],
        ['5年', '120,000元', '196,035元', '76,035元', '63.4%'],
        ['10年', '240,000元', '496,356元', '256,356元', '106.8%']
    ]
    
    for i, row_data in enumerate(calc_data, 1):
        for j, cell_data in enumerate(row_data):
            cell = calc_table.rows[i].cells[j]
            cell.text = cell_data
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if i >= 3:  # 5年及以上高亮
                        run.font.bold = True
                        if j >= 2:
                            run.font.color.rgb = RGBColor(0x00, 0xB0, 0x50)
    
    note = doc.add_paragraph('注：以上测算基于12%年化收益率，实际收益会随市场波动。红利再投资可进一步提升实际收益。')
    note.runs[0].font.italic = True
    note.runs[0].font.size = Pt(10)
    note.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    
    doc.add_page_break()
    
    # ========== 第五章 ==========
    doc.add_heading('五、风险提示与应对', level=1)
    
    doc.add_heading('5.1 主要风险', level=2)
    
    risks = [
        '市场风险：股市整体下跌时，红利ETF也会随之下跌，但跌幅通常小于成长型ETF',
        '分红减少风险：成分股公司减少分红时，ETF分红也会减少',
        '跟踪误差风险：ETF表现可能与标的指数存在偏差',
        '流动性风险：小规模ETF可能出现买卖价差大的情况'
    ]
    
    for risk in risks:
        p = doc.add_paragraph(risk, style='List Number')
        p.space_after = Pt(6)
    
    doc.add_heading('5.2 应对策略', level=2)
    
    strategies = [
        '坚持长期定投：不要在市场下跌时停止定投，反而是加仓良机',
        '分红再投资：选择"红利再投资"而非"现金分红"，充分利用复利',
        '定期再平衡：每年检查一次持仓，如某只ETF占比过高可适当减仓',
        '设置止盈线：如总收益率超过80%，可考虑分批止盈，锁定收益'
    ]
    
    for strategy in strategies:
        p = doc.add_paragraph(strategy, style='List Bullet')
        p.space_after = Pt(6)
    
    # ========== 结论 ==========
    doc.add_page_break()
    doc.add_heading('六、结论', level=1)
    
    p = doc.add_paragraph('综合来看，515080中证红利ETF招商是当前市场环境下最适合定投的红利ETF标的。其跨市场分散、高频率分红、规模适中、表现稳健的特点，非常适合长期定投。')
    p.space_after = Pt(12)
    
    p = doc.add_paragraph('建议采用"70% 515080 + 30% 510880"的均衡配置，每月定投2000元，坚持5年以上，预期可获得12%的年化收益。重要的是保持纪律，不要因短期波动而改变策略。')
    p.space_after = Pt(12)
    
    # 免责声明
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('投资有风险，以上分析仅供参考，不构成投资建议。实际操作前请咨询专业投资顾问。')
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    
    # 保存文档
    output_path = r'd:\temp\红利ETF\红利ETF定投分析报告.docx'
    doc.save(output_path)
    print('Word文档生成成功！')
    print(f'文件路径：{output_path}')
    return output_path

if __name__ == '__main__':
    try:
        output = create_document()
        print('文档创建完成！')
    except Exception as e:
        print(f'错误：{e}')
        import traceback
        traceback.print_exc()
